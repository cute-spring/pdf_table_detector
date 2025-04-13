import sys
import json
import re
from docx import Document
from docx.shared import Pt
from docx.table import _Cell, Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
import logging

# --- Configuration ---
DEFAULT_OLLAMA_URL = "http://localhost:11434/api/generate"
DEFAULT_OLLAMA_MODEL = "phi4" # Or your preferred model
DEFAULT_TABLE_STYLE_FALLBACK = 'Table Normal' # Usually safe fallback

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Helper Functions ---

def extract_table_to_textual_representation(table: Table):
    """
    Extracts table data into a simple textual representation.
    Basic handling for merged cells (outputs content from the top-left cell).
    """
    text_repr = []
    logging.info(f"Extracting table with {len(table.rows)} rows and {len(table.columns)} columns.")
    try:
        # Access the underlying grid of cells (_cells)
        grid = table._cells
        num_rows = len(table.rows)
        num_cols = len(table.columns)

        for i in range(num_rows):
            row_cells = []
            for j in range(num_cols):
                try:
                    # Calculate the flat index for _cells
                    cell_idx = i * num_cols + j
                    if cell_idx < len(grid):
                        cell = grid[cell_idx]
                        # Replace newline characters for better text representation
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_cells.append(cell_text)
                    else:
                         # Should not happen with table._cells if table structure is standard
                         logging.warning(f"Calculated cell index {cell_idx} out of bounds for grid length {len(grid)} at ({i},{j})")
                         row_cells.append("(cell-missing?)")
                except Exception as cell_e:
                    logging.error(f"Error processing cell at ({i},{j}): {cell_e}")
                    row_cells.append("(error reading cell)")
            text_repr.append(" | ".join(row_cells))

    except Exception as e:
        logging.error(f"Error extracting table content: {e}")
        return None # Indicate failure

    return "\n".join(text_repr)


def simplify_table_with_ollama(table_text,
                              model_name=DEFAULT_OLLAMA_MODEL,
                              ollama_url=DEFAULT_OLLAMA_URL,
                              timeout=120):
    """
    Sends the extracted table text to the Ollama LLM for simplification
    (handling merged cells by filling data) and returns a Markdown table.
    """
    logging.info(f"Sending table data to Ollama model '{model_name}' at {ollama_url}")

    prompt = f"""
You are an expert data processor. Below is a textual representation of a table extracted from a document.
The table may contain merged cells. In the representation, cells that were part of a vertical or horizontal merge might contain repeated text identical to the top-left cell of the merge area. Some might appear as empty strings if the source cell was empty.

Your task is to:
1. Analyze the structure and content of the table representation.
2. Intelligently "unmerge" all cells. Identify regions with repeated text (or potentially adjacent empty cells following text) that indicate a merge.
3. Fill in all cells that were part of an original merge by replicating the data from the conceptual top-left cell of that merged region. Ensure the text is consistent downwards for vertical merges and rightwards for horizontal merges.
4. Ensure the final table has a consistent number of columns in every row and conceptually no merged cells.
5. Output ONLY the resulting table content as a standard Markdown table. Do not include any explanations, introductions, or summaries before or after the Markdown table itself.

Original table representation:
--- START TABLE ---
{table_text}
--- END TABLE ---

Unmerged Markdown table:
"""

    data = {
        "model": model_name,
        "prompt": prompt,
        "stream": False
    }

    try:
        response = requests.post(ollama_url, json=data, timeout=timeout)
        response.raise_for_status() # Checks for HTTP errors (4xx, 5xx)

        result = response.json()
        simplified_table_md = result.get('response', '').strip()

        if not simplified_table_md:
            logging.error("LLM response was empty.")
            return None

        logging.info("Received non-empty response from Ollama.")

        # Basic check for Markdown structure
        if not simplified_table_md.startswith('|') or simplified_table_md.count('\n') < 1:
             logging.warning("LLM response might not be a valid Markdown table. Output starts with: %s", simplified_table_md[:100])
             logging.debug(f"LLM Raw Response:\n{simplified_table_md}")

        return simplified_table_md

    except requests.exceptions.Timeout:
        logging.error(f"Request to Ollama timed out after {timeout} seconds.")
        return None
    except requests.exceptions.ConnectionError:
        logging.error(f"Could not connect to Ollama at {ollama_url}. Is it running?")
        return None
    except requests.exceptions.RequestException as e: # Includes HTTPError from raise_for_status
        logging.error(f"Error communicating with Ollama: {e}")
        if hasattr(e, 'response') and e.response is not None:
            logging.error(f"Ollama response status: {e.response.status_code}")
            logging.error(f"Ollama response text: {e.response.text}")
        return None
    except json.JSONDecodeError as e:
        logging.error(f"Error decoding JSON response from Ollama: {e}")
        logging.debug(f"Raw response content: {response.text if 'response' in locals() else 'Response object not available'}")
        return None
    except Exception as e:
        logging.error(f"An unexpected error occurred during LLM interaction: {e}")
        return None


def parse_markdown_table(markdown_string):
    """
    Parses a Markdown table string into a list of lists (rows and cells).
    Improved robustness for various Markdown table formats.
    """
    logging.info("Parsing simplified Markdown table from LLM response.")
    lines = markdown_string.strip().split('\n')
    data = []
    header_separator_found = False
    num_cols = 0

    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|') or not line.endswith('|'):
            # Skip lines not conforming to basic |...| structure, unless it's the separator
            # Check for separator even if it doesn't start/end with '|' sometimes LLMs output slightly malformed separators
            if re.match(r"^\s*\|?\s*[-:|]+\s*\|?(\s*[-:|]+\s*\|?\s*)*$", line):
                 pass # Let the separator check below handle it
            else:
                 logging.debug(f"Skipping line {i+1} (doesn't seem like table data or separator): {line}")
                 continue

        # Check for header separator line (e.g., |---|---| or :---|:---: or ---|--- )
        if re.match(r"^\s*\|?\s*[-:|]+\s*\|?(\s*[-:|]+\s*\|?\s*)*$", line):
            if not header_separator_found:
                 logging.debug(f"Found header separator at line {i+1}: {line}")
                 header_separator_found = True
                 # Determine column count from the header row just processed (if available)
                 if data:
                     num_cols = len(data[-1])
                     logging.info(f"Determined {num_cols} columns based on header row.")
                 else:
                     # Try inferring from separator itself (less reliable)
                     inferred_cols = line.count('|') - 1
                     if inferred_cols > 0:
                         num_cols = inferred_cols
                         logging.warning(f"No header row found before separator, inferring {num_cols} columns from separator line. This might be inaccurate.")
                     else:
                         logging.warning(f"Could not determine column count from header or separator at line {i+1}.")
                         # Will try to determine from the first actual data row later
            continue # Skip the separator line itself from data rows

        # Extract cells from the line
        cells = [cell.strip() for cell in line.strip('|').split('|')]

        # Determine num_cols from the first valid row (header or data) if not already set
        if num_cols == 0 and len(cells) > 0:
             num_cols = len(cells)
             logging.info(f"Setting expected column count to {num_cols} based on first row (line {i+1}).")

        # Append row data, ensuring consistency if num_cols is known
        if num_cols > 0:
            if len(cells) != num_cols:
                logging.warning(f"Inconsistent column count at line {i+1}. Expected {num_cols}, found {len(cells)}. Adjusting row.")
                # Pad or truncate
                if len(cells) < num_cols:
                    cells.extend([""] * (num_cols - len(cells)))
                else:
                    cells = cells[:num_cols]
            data.append(cells)
        elif len(cells) > 0: # Only append if it looks like a valid row, even if num_cols unknown
             data.append(cells) # Add first row before num_cols is finalized


    if not data:
        logging.warning("Could not parse any data rows from the Markdown.")
        return None
    if num_cols == 0 and data: # Edge case: only one row parsed, num_cols never finalized from separator or later rows
         num_cols = len(data[0])
         logging.info(f"Finalized column count to {num_cols} based on the single row found.")

    # Basic validation: Ensure all rows have the finalized num_cols if determined
    if num_cols > 0:
        for i_row, row in enumerate(data):
            if len(row) != num_cols:
                logging.warning(f"Post-processing found row {i_row} has {len(row)} cols, expected {num_cols}. Adjusting.")
                if len(row) < num_cols:
                    row.extend([""] * (num_cols - len(row)))
                else:
                    data[i_row] = row[:num_cols]

    logging.info(f"Successfully parsed {len(data)} rows and {num_cols} columns from Markdown.")
    return data


# ***************************************************************************
# *** Function below was modified to enhance logging and error handling ***
# ***************************************************************************
def insert_table_after(document, ref_table_element, data, original_style_name=None):
    """
    Inserts a new table populated with data immediately after the reference table element.
    Applies the `original_style_name` if provided and valid, otherwise uses fallbacks.
    Includes enhanced logging and error handling for insertion issues.
    """
    if not data or not data[0]:
        logging.error("No data provided to create the new table.")
        return None

    num_rows = len(data)
    num_cols = len(data[0])
    if num_cols == 0:
         logging.error("Cannot create a table with 0 columns based on parsed data.")
         return None
    logging.info(f"Preparing to create a new table with {num_rows} rows and {num_cols} columns.")

    new_table = None
    applied_style = None
    new_table_element = None # Initialize here for potential cleanup later

    # --- Try creating the table with styles ---
    try:
        # 1. Attempt to use the original table's style name
        if original_style_name:
            try:
                logging.info(f"Attempting to apply original style: '{original_style_name}'")
                safe_rows = max(1, num_rows)
                safe_cols = max(1, num_cols)
                new_table = document.add_table(rows=safe_rows, cols=safe_cols, style=original_style_name)
                applied_style = original_style_name
                logging.info(f"Successfully created table object with original style '{applied_style}'.")
                if num_rows < 1: logging.warning("Requested 0 rows, created table with 1 row.")
            except KeyError:
                logging.warning(f"Original style '{original_style_name}' not found. Proceeding to fallback.")
                new_table = None
            except Exception as e:
                logging.error(f"Error applying original style '{original_style_name}': {e}. Proceeding to fallback.")
                new_table = None

        # 2. Fallback to DEFAULT_TABLE_STYLE_FALLBACK
        if new_table is None:
            try:
                logging.info(f"Attempting fallback style: '{DEFAULT_TABLE_STYLE_FALLBACK}'")
                safe_rows = max(1, num_rows)
                safe_cols = max(1, num_cols)
                new_table = document.add_table(rows=safe_rows, cols=safe_cols, style=DEFAULT_TABLE_STYLE_FALLBACK)
                applied_style = DEFAULT_TABLE_STYLE_FALLBACK
                logging.info(f"Successfully created table object with fallback style '{applied_style}'.")
                if num_rows < 1: logging.warning("Requested 0 rows, created table with 1 row.")
            except KeyError:
                logging.warning(f"Fallback style '{DEFAULT_TABLE_STYLE_FALLBACK}' not found. Proceeding to default.")
                new_table = None
            except Exception as e:
                logging.error(f"Error applying fallback style '{DEFAULT_TABLE_STYLE_FALLBACK}': {e}. Proceeding to default.")
                new_table = None

        # 3. Fallback to document's default style
        if new_table is None:
            try:
                logging.info("Attempting to create table object with document's default style.")
                safe_rows = max(1, num_rows)
                safe_cols = max(1, num_cols)
                new_table = document.add_table(rows=safe_rows, cols=safe_cols) # Omit style argument
                applied_style = "[Document Default]"
                logging.info(f"Successfully created table object with {applied_style}.")
                if num_rows < 1: logging.warning("Requested 0 rows, created table with 1 row.")
            except Exception as e:
                logging.error(f"Failed to create table object even with default style: {e}", exc_info=True)
                return None # Critical failure

    except Exception as table_creation_e:
        logging.error(f"Unexpected error during table object creation phase: {table_creation_e}", exc_info=True)
        return None


    # --- Populate and Insert the Table ---
    if new_table is None:
        logging.error("Table creation variable 'new_table' is None before population/insertion. Should not happen if creation succeeded.")
        return None

    # Store the OXML element early for potential cleanup later
    new_table_element = new_table._element

    try: # Wrap population and insertion in a larger try block for better cleanup
        new_table.autofit = True

        # --- Populate Table Data ---
        if not hasattr(new_table, 'rows'):
            logging.error("Newly created table object unexpectedly lacks 'rows' attribute. Cannot populate.")
            raise RuntimeError("Table created without rows attribute") # Raise to trigger cleanup

        if num_rows == 0 and len(new_table.rows) > 0:
            logging.info("Clearing the placeholder row created for a 0-row request.")
            placeholder_row = new_table.rows[0]
            for cell_obj in placeholder_row.cells:
                 for para in list(cell_obj.paragraphs): # Use list() for safe iteration
                    p_element = para._element
                    if p_element.getparent() is not None: p_element.getparent().remove(p_element)
                 cell_obj.add_paragraph("") # Add single empty paragraph

        elif num_rows > 0:
            logging.info("Populating new table data.")
            for i, row_data in enumerate(data):
                if i >= len(new_table.rows):
                    logging.warning(f"Data row index {i} exceeds available rows ({len(new_table.rows)}). Stopping population.")
                    break
                row_cells = new_table.rows[i].cells
                if not row_cells:
                    logging.warning(f"Row {i} in created table has no cells. Skipping population for this row.")
                    continue
                for j, cell_text in enumerate(row_data):
                    if j < len(row_cells):
                        cleaned_text = re.sub(r'[*_`]', '', str(cell_text))
                        cell_obj = row_cells[j]
                        # Clear existing paragraphs first
                        for para in list(cell_obj.paragraphs):
                             p_element = para._element
                             if p_element.getparent() is not None: p_element.getparent().remove(p_element)
                        cell_obj.add_paragraph(cleaned_text)
                    else:
                        logging.warning(f"Data column index {j} exceeds table column count {len(row_cells)} in row {i}. Skipping cell.")
            logging.info("Finished populating table data (initial attempt).")


        # --- Insert the new table immediately after the reference table element using OXML ---
        logging.debug("Attempting OXML insertion...")
        original_table_element = ref_table_element

        parent_element = original_table_element.getparent()
        if parent_element is None:
            logging.error("Could not find the parent element of the original table in OXML. Cannot insert correctly.")
            raise RuntimeError("Original table has no parent element") # Raise to trigger cleanup

        logging.debug(f"Original table parent element: {parent_element.tag}")

        try:
            original_table_index = parent_element.index(original_table_element)
            logging.debug(f"Found original table at index {original_table_index} within parent.")
        except ValueError:
            logging.error("Could not find the original table element within its parent during insertion. Document structure might be complex or changed.", exc_info=True)
            raise # Re-raise ValueError to trigger cleanup

        # --- Core Insertion Logic ---
        logging.debug(f"Inserting new table element {new_table_element.tag} at index {original_table_index + 1}")
        parent_element.insert(original_table_index + 1, new_table_element)
        logging.info(f"Successfully moved new table ({applied_style}) OXML element after the original table.")

        # Verify insertion (optional but good for debug)
        try:
            check_idx = parent_element.index(new_table_element)
            if check_idx != original_table_index + 1:
                 logging.warning(f"Verification: New table is at index {check_idx}, expected {original_table_index + 1}. Potential structure issue.")
            else:
                 logging.debug(f"Verified: New table is now at expected index {check_idx} in parent.")
        except ValueError:
            logging.error("CRITICAL: New table element not found in parent immediately after insertion! Insertion likely failed silently.")
            # This is a serious issue if it occurs.
            raise RuntimeError("Failed to verify new table element position after insertion.")

        # Add a blank paragraph between the tables for visual spacing
        logging.debug("Adding paragraph separator...")
        p = OxmlElement("w:p")
        # Insert paragraph *before* the new table (which is now at original_table_index + 1)
        # So, insert paragraph also at original_table_index + 1
        parent_element.insert(original_table_index + 1, p)
        logging.debug("Inserted paragraph separator between tables.")
        # --- End of Core Insertion Logic ---

        logging.info("Successfully completed OXML insertion steps.")
        return new_table # Return the created and inserted table object

    except Exception as e_insert_populate:
        # Catch ANY exception during population or insertion
        logging.error(f"Error occurred during table population or OXML insertion: {e_insert_populate}", exc_info=True) # Log traceback

        # Attempt to clean up the table that was added (likely at the end initially)
        logging.warning("Attempting to clean up partially created/inserted table due to error.")
        # Use the stored element. Check if it was created.
        if new_table_element is not None:
            current_parent = new_table_element.getparent()
            if current_parent is not None:
                try:
                    current_parent.remove(new_table_element)
                    logging.info("Successfully removed potentially orphaned new table element during cleanup.")
                except Exception as cleanup_e:
                    logging.error(f"Error during table cleanup after insertion failure: {cleanup_e}", exc_info=True)
            else:
                logging.warning("Could not find parent of the new table element during cleanup (might already be detached or structure is broken).")
        else:
            logging.warning("New table element was not created/available for cleanup.")

        return None # Indicate failure


# --- Main Workflow ---
def process_docx(input_path, output_path, ollama_model=DEFAULT_OLLAMA_MODEL, ollama_url=DEFAULT_OLLAMA_URL):
    """
    Processes the DOCX: extracts each table, simplifies via LLM,
    re-inserts the simplified table using the original's style (with fallbacks).
    Includes enhanced logging for insertion success/failure tracking.
    """
    logging.info(f"Starting processing for DOCX: '{input_path}'")
    try:
        document = Document(input_path)
    except Exception as e:
        logging.error(f"Failed to open DOCX file '{input_path}': {e}")
        return

    if not document.tables:
        logging.info("No tables found in the document.")
        return

    logging.info(f"Found {len(document.tables)} tables in the document. Processing all.")
    tables_processed_count = 0
    tables_failed_count = 0

    # Create a static list of table info (object + OXML element) to iterate over
    original_tables_info = []
    for table in document.tables:
        if table is not None and hasattr(table, '_element'):
            original_tables_info.append({'table_obj': table, 'element': table._element})
        else:
             logging.warning("Found an invalid table object in the document's initial list, skipping.")

    logging.info(f"Prepared list of {len(original_tables_info)} valid tables to process.")

    for i, table_info in enumerate(original_tables_info):
        original_table = table_info['table_obj']
        original_table_element = table_info['element']

        logging.info(f"--- Processing Table {i+1} of {len(original_tables_info)} ---")

        # --- Get Original Style Name ---
        original_style_name = None
        try:
            if hasattr(original_table, 'style') and original_table.style and hasattr(original_table.style, 'name'):
                style_name_candidate = original_table.style.name
                if style_name_candidate:
                    normalized_style_name = style_name_candidate.lower().strip()
                    # Common default styles to treat as 'no specific style'
                    default_styles = ["none", "table normal", "normal table", "grid table"]
                    if normalized_style_name in default_styles:
                        logging.info(f"Original table uses default-like style ('{style_name_candidate}'). Will use preferred fallback '{DEFAULT_TABLE_STYLE_FALLBACK}'.")
                        original_style_name = None
                    else:
                        original_style_name = style_name_candidate
                        logging.info(f"Detected specific original table style: '{original_style_name}'")
                else:
                    logging.info("Original table style name is missing or empty. Using fallback logic.")
                    original_style_name = None
            else:
                logging.info("Original table lacks specific named style attribute. Using fallback logic.")
                original_style_name = None
        except AttributeError:
             logging.warning("Error accessing style attributes. Using fallback logic.")
             original_style_name = None
        except Exception as e:
            logging.warning(f"Error determining style name for table {i+1}: {e}. Using fallback logic.", exc_info=True)
            original_style_name = None


        # 1. Extract Table Content
        logging.debug("Step 1: Extracting table content...")
        table_text_repr = extract_table_to_textual_representation(original_table)
        if table_text_repr is None:
            logging.error(f"Failed Step 1 (Extract) for table {i+1}. Skipping.")
            tables_failed_count += 1
            continue

        logging.debug(f"Extracted Text Representation (Table {i+1}):\n{table_text_repr[:500]}...") # Log first 500 chars

        # 2. Simplify Table via LLM
        logging.debug("Step 2: Sending table to LLM for simplification...")
        simplified_markdown = simplify_table_with_ollama(table_text_repr, ollama_model, ollama_url)
        if not simplified_markdown:
            logging.error(f"Failed Step 2 (LLM) for table {i+1}. Skipping.")
            tables_failed_count += 1
            continue

        logging.debug(f"Simplified Markdown received (Table {i+1}):\n{simplified_markdown[:500]}...") # Log first 500 chars

        # 3. Parse LLM Response
        logging.debug("Step 3: Parsing simplified Markdown...")
        simplified_data = parse_markdown_table(simplified_markdown)
        if not simplified_data:
            logging.error(f"Failed Step 3 (Parse) for table {i+1}. Skipping.")
            tables_failed_count += 1
            continue
        logging.info(f"Parsed simplified data: {len(simplified_data)} rows, {len(simplified_data[0]) if simplified_data else 0} cols.")


        # ***************************************************************************
        # *** Code below logs result of insert_table_after more clearly         ***
        # ***************************************************************************
        # 4. Insert Simplified Table
        logging.info(f"Step 4: Calling insert_table_after for original table {i+1}")
        inserted_table_obj = insert_table_after( # Capture the return value
            document,
            original_table_element,
            simplified_data,
            original_style_name
        )

        # --- Log the result ---
        if inserted_table_obj is None:
            logging.error(f"insert_table_after FAILED for original table {i+1}. No table was inserted or insertion process failed.")
            tables_failed_count += 1 # Count insertion failure as a failure
        else:
            logging.info(f"insert_table_after SUCCEEDED for original table {i+1}. Returned object is type: {type(inserted_table_obj)}")
            # Perform a quick check if the object is indeed a Table
            if isinstance(inserted_table_obj, Table):
                 logging.info("Returned object is a valid Table instance.")
                 tables_processed_count += 1
            else:
                 # This case should theoretically not happen if insert_table_after logic is correct
                 logging.error(f"insert_table_after returned an unexpected object type: {type(inserted_table_obj)}. Counting as failed.")
                 tables_failed_count += 1
        # --- End logging result ---


        logging.info(f"--- Finished Processing Table {i+1} ---")
        if i < len(original_tables_info) - 1:
            logging.info("-" * 30) # Visual separator

    # --- Save Final Document ---
    logging.info(f"Finished processing all tables. Processed successfully: {tables_processed_count}, Failed/Skipped: {tables_failed_count}.")
    try:
        logging.info(f"Attempting to save document to '{output_path}'...")
        document.save(output_path)
        logging.info(f"Successfully saved updated document to '{output_path}'")
    except Exception as e:
        logging.error(f"Failed to save the final updated DOCX file '{output_path}': {e}", exc_info=True)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Simplify complex tables in DOCX using Ollama LLM, reusing original table styles.")
    parser.add_argument("input_docx", help="Path to the input DOCX file.")
    parser.add_argument("output_docx", help="Path to save the modified DOCX file.")
    parser.add_argument("--model", default=DEFAULT_OLLAMA_MODEL, help=f"Ollama model name (default: {DEFAULT_OLLAMA_MODEL}).")
    parser.add_argument("--url", default=DEFAULT_OLLAMA_URL, help=f"Ollama API URL (default: {DEFAULT_OLLAMA_URL}).")
    parser.add_argument("--debug", action='store_true', help="Enable debug logging.")

    args = parser.parse_args()

    # Adjust logging level based on debug flag BEFORE processing
    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
        logging.debug("Debug logging enabled.")
    else:
        logging.getLogger().setLevel(logging.INFO) # Ensure INFO level if not debug

    process_docx(
        input_path=args.input_docx,
        output_path=args.output_docx,
        ollama_model=args.model,
        ollama_url=args.url
    )